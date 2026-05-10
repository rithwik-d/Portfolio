from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


OUT_DIR = Path(__file__).resolve().parent
BASENAME = "Rithwik_Reddy_ATS_Resume"


resume = {
    "name": "Rithwik Reddy Donthi Reddy",
    "title": "Software Engineer | Full-Stack Development | AI/ML",
    "contact": (
        "Norman, OK | 405-618-0095 | rithwik6196@gmail.com | "
        "linkedin.com/in/rithwik-reddyd | github.com/rithwik-d | "
        "rithwik-d.github.io/Portfolio/"
    ),
    "sections": [
        {
            "heading": "Professional Summary",
            "items": [
                (
                    "Computer Science M.S. graduate from the University of Oklahoma with hands-on "
                    "full-stack software engineering and AI/ML project experience. Built and deployed "
                    "applications using React, Node.js, Express.js, PostgreSQL, MongoDB, authentication, "
                    "REST APIs, CI/CD, and cloud hosting."
                )
            ],
        },
        {
            "heading": "Technical Skills",
            "items": [
                "Languages: JavaScript, Python, C++, HTML5, CSS3, SQL",
                "Frontend: React, React Router, EJS, Bootstrap, jQuery, Responsive Design",
                "Backend: Node.js, Express.js, REST APIs, Authentication, Sessions, bcrypt, Nodemailer",
                "Databases: PostgreSQL, MongoDB, Mongoose, Microsoft SQL Server",
                "AI/ML: Machine Learning, Computer Vision, YOLOv5, PyTorch, OpenCV, NumPy, Pandas, scikit-learn, TensorBoard",
                "Cloud/Tools: Google Cloud Run, Cloud SQL, Vercel, GitHub Pages, GitHub Actions, Git, Postman, VS Code, Figma, Azure",
            ],
        },
        {
            "heading": "Education",
            "items": [
                "University of Oklahoma, Norman, OK | M.S. in Computer Science | GPA: 3.67/4.0 | May 2026",
                "VNR VJIET, Hyderabad, India | B.S. in Computer Science | GPA: 8.4/10 | Mar 2024",
            ],
        },
        {
            "heading": "Projects",
            "projects": [
                {
                    "name": "Inkline Journal - Full-Stack Blogging Platform",
                    "stack": "Node.js, Express.js, EJS, PostgreSQL, bcrypt, Nodemailer",
                    "links": "Live: https://www.inklinejournal.com | GitHub: github.com/rithwik-d/inkline-journal",
                    "bullets": [
                        "Built and deployed a multi-user blogging platform with public feed, signup/login, author dashboard, and post ownership controls.",
                        "Implemented authentication with bcrypt password hashing, express-session, one-time email verification tokens, verified-only sign-in, and SMTP email delivery.",
                        "Designed PostgreSQL persistence with users, posts, verification tokens, indexed queries, and full CRUD workflows for publish, unpublish, edit, and delete actions.",
                        "Configured production deployment with custom domain/SSL, cloud database settings, and environment-driven application configuration.",
                    ],
                },
                {
                    "name": "IPMS - MERN Application",
                    "stack": "React, React Router, Node.js, Express.js, MongoDB, Mongoose, Axios, GitHub Actions",
                    "links": "GitHub: github.com/rithwik-d/IPMS",
                    "bullets": [
                        "Developed a full-stack app with React client, Express REST API, MongoDB/Mongoose models, CORS, and environment-based configuration.",
                        "Built user creation and email routes with JSON APIs, Nodemailer service layer, and Axios-driven client/server communication.",
                        "Added GitHub Actions CI workflow using Node.js 18 and MongoDB service container to install dependencies and validate React production builds.",
                    ],
                },
                {
                    "name": "Helmet and License Plate Detection - Computer Vision",
                    "stack": "Python, YOLOv5, PyTorch, OpenCV, Pandas",
                    "links": "GitHub: github.com/rithwik-d/helmet-and-licence-plate-detection",
                    "bullets": [
                        "Trained a YOLOv5 object detection model to identify helmet usage and license plates in rider images/videos.",
                        "Built detection workflow to flag non-helmet riders, crop license plates, save plate text/images, and log results with timestamps to CSV.",
                        "Used PyTorch, OpenCV, NumPy, Pandas, TensorBoard, and model-export tooling to support training, detection, and result analysis.",
                    ],
                },
            ],
        },
        {
            "heading": "Leadership",
            "items": [
                (
                    "Founding Member and Social Media Head, Krithomedh | Helped launch a technical "
                    "community and led outreach for events/workshops, strengthening collaboration, "
                    "event planning, and communication."
                )
            ],
        },
    ],
}


def markdown_text() -> str:
    lines = [
        f"# {resume['name']}",
        f"**{resume['title']}**",
        resume["contact"],
        "",
    ]
    for section in resume["sections"]:
        lines.append(f"## {section['heading']}")
        if "projects" in section:
            for project in section["projects"]:
                lines.append("")
                lines.append(f"**{project['name']} | {project['stack']}**")
                lines.append(project["links"])
                for bullet in project["bullets"]:
                    lines.append(f"- {bullet}")
        else:
            for item in section["items"]:
                if section["heading"] in {"Technical Skills"}:
                    lines.append(f"- {item}")
                else:
                    lines.append(item)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def plain_text() -> str:
    lines = []
    for line in markdown_text().splitlines():
        if line.startswith("## "):
            lines.append(line[3:])
        elif line.startswith("# "):
            lines.append(line[2:])
        elif line.startswith("- "):
            lines.append(f"  - {line[2:]}")
        else:
            lines.append(line.replace("**", ""))
    return "\n".join(lines) + "\n"


def set_run_font(paragraph, size=10, bold=False):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold


def add_docx_paragraph(document, text, size=9.6, bold=False, align=None, space_after=0):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    return paragraph


def add_docx_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.first_line_indent = Inches(-0.1)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(8.8)
    return paragraph


def build_docx():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9.4)

    add_docx_paragraph(document, resume["name"], size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_docx_paragraph(document, resume["title"], size=10.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_docx_paragraph(document, resume["contact"], size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    for section_data in resume["sections"]:
        add_docx_paragraph(document, section_data["heading"].upper(), size=10.2, bold=True, space_after=1)
        if "projects" in section_data:
            for project in section_data["projects"]:
                add_docx_paragraph(document, f"{project['name']} | {project['stack']}", size=9.4, bold=True)
                add_docx_paragraph(document, project["links"], size=8.4)
                for bullet in project["bullets"]:
                    add_docx_bullet(document, bullet)
        elif section_data["heading"] == "Technical Skills":
            for item in section_data["items"]:
                add_docx_bullet(document, item)
        else:
            for item in section_data["items"]:
                add_docx_paragraph(document, item, size=9.1)

    document.save(OUT_DIR / f"{BASENAME}.docx")


def pdf_para_style(name, font_size, leading, bold=False, alignment=TA_LEFT, space_after=0):
    return ParagraphStyle(
        name,
        fontName="Helvetica-Bold" if bold else "Helvetica",
        fontSize=font_size,
        leading=leading,
        alignment=alignment,
        textColor=colors.black,
        spaceAfter=space_after,
    )


def build_pdf():
    path = OUT_DIR / f"{BASENAME}.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=0.48 * inch,
        leftMargin=0.48 * inch,
        topMargin=0.42 * inch,
        bottomMargin=0.42 * inch,
    )

    styles = getSampleStyleSheet()
    styles.add(pdf_para_style("ResumeName", 15.5, 17, bold=True, alignment=TA_CENTER))
    styles.add(pdf_para_style("ResumeTitle", 9.6, 11, bold=True, alignment=TA_CENTER))
    styles.add(pdf_para_style("Contact", 7.9, 9.4, alignment=TA_CENTER, space_after=4))
    styles.add(pdf_para_style("Section", 9.8, 11, bold=True, space_after=1))
    styles.add(pdf_para_style("Body", 8.55, 10.2, space_after=1))
    styles.add(pdf_para_style("Project", 8.8, 10.5, bold=True))
    styles.add(pdf_para_style("Link", 7.9, 9.4))
    styles.add(
        ParagraphStyle(
            "ResumeBullet",
            fontName="Helvetica",
            fontSize=8.2,
            leading=9.5,
            leftIndent=12,
            firstLineIndent=-7,
            bulletIndent=3,
            spaceAfter=0.4,
        )
    )

    story = [
        Paragraph(escape(resume["name"]), styles["ResumeName"]),
        Paragraph(escape(resume["title"]), styles["ResumeTitle"]),
        Paragraph(escape(resume["contact"]), styles["Contact"]),
    ]

    for section_data in resume["sections"]:
        story.append(Spacer(1, 2.2))
        story.append(Paragraph(escape(section_data["heading"].upper()), styles["Section"]))
        if "projects" in section_data:
            for project in section_data["projects"]:
                story.append(
                    Paragraph(
                        f"<b>{escape(project['name'])}</b> | {escape(project['stack'])}",
                        styles["Project"],
                    )
                )
                story.append(Paragraph(escape(project["links"]), styles["Link"]))
                for bullet in project["bullets"]:
                    story.append(Paragraph(escape(bullet), styles["ResumeBullet"], bulletText="-"))
        elif section_data["heading"] == "Technical Skills":
            for item in section_data["items"]:
                story.append(Paragraph(escape(item), styles["ResumeBullet"], bulletText="-"))
        else:
            for item in section_data["items"]:
                story.append(Paragraph(escape(item), styles["Body"]))

    doc.build(story)


def main():
    (OUT_DIR / f"{BASENAME}.md").write_text(markdown_text(), encoding="utf-8")
    (OUT_DIR / f"{BASENAME}.txt").write_text(plain_text(), encoding="utf-8")
    build_docx()
    build_pdf()


if __name__ == "__main__":
    main()
